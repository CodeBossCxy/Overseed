from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT = "Overseed_Infrastructure_Survival_Budget.docx"
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
GRAY = "F2F4F7"
MUTED = RGBColor(89, 99, 110)
WHITE = RGBColor(255, 255, 255)
INK = RGBColor(31, 41, 55)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_run(run, size=9, bold=False, color=INK, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_text(cell, text, size=8.3, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    format_run(p.add_run(text), size=size, bold=bold, color=color)


def add_budget_table(doc, rows):
    headers = ["Infrastructure", "Cost driver / billing basis", "Units / month", "Unit price", "Monthly forecast", "Annualized forecast", "Essential?", "Notes / pricing source"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [1900, 2300, 1050, 1050, 1250, 1450, 950, 2450]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        values = [row[0], row[1], "", "", "", "", row[2], row[3]]
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (2, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
            if idx % 2 == 1:
                shade(cells[i], PALE_BLUE)
    set_fixed_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_formula_note(doc):
    p = doc.add_paragraph(style="Budget Note")
    r = p.add_run("Calculation rule: ")
    format_run(r, size=9, bold=True, color=RGBColor(31, 78, 120))
    format_run(p.add_run("Monthly forecast = units per month × unit price. Annualized forecast = monthly forecast × 12. For annual flat fees, enter 1/12 of the fee as the monthly forecast."), size=9)


def add_section_table(doc, title, intro, rows):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(intro)
    p.style = doc.styles["Normal"]
    add_formula_note(doc)
    add_budget_table(doc, rows)


def add_summary_table(doc):
    headers = ["Budget layer", "Monthly total", "Annual total", "Cash runway (months)", "Purpose"]
    rows = [
        ("Bare-minimum survival", "", "", "", "Only services required to keep the product available and secure."),
        ("Current operating baseline", "", "", "", "Survival layer plus normal production usage."),
        ("Growth-ready plan", "", "", "", "Baseline plus optional capacity, discovery, and experimental services."),
        ("Contingency reserve", "", "", "", "Recommended: 10%–20% of annual operating infrastructure."),
        ("Total funding target", "", "", "", "Annual plan plus contingency and any one-time setup costs."),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [2350, 1450, 1450, 1550, 5550]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
            if idx % 2 == 1:
                shade(cells[i], PALE_BLUE)
    set_fixed_table(table, widths)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08
for level, size, before, after in ((1, 16, 13, 6), (2, 12, 9, 4)):
    style = styles[f"Heading {level}"]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(31, 78, 120)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

note_style = styles.add_style("Budget Note", WD_STYLE_TYPE.PARAGRAPH)
note_style.font.name = "Aptos"
note_style.font.size = Pt(9)
note_style.paragraph_format.space_before = Pt(2)
note_style.paragraph_format.space_after = Pt(6)
note_style.paragraph_format.left_indent = Inches(0.12)
note_style.paragraph_format.right_indent = Inches(0.12)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
format_run(header.add_run("OVERSEED  |  INFRASTRUCTURE COST PLANNING"), size=8, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_run(footer.add_run("Working budget template • Update assumptions whenever vendor pricing or usage changes"), size=8, color=MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
format_run(p.add_run("OVERSEED"), size=10, bold=True, color=RGBColor(31, 78, 120))
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
format_run(p.add_run("Infrastructure Survival Budget"), size=25, bold=True, color=RGBColor(20, 45, 70), font="Aptos Display")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
format_run(p.add_run("Editable worksheet for unit pricing, monthly cost, annualized forecasts, and runway planning"), size=11, color=MUTED)

intro = doc.add_paragraph(style="Budget Note")
format_run(intro.add_run("How to use this document: "), size=9.5, bold=True, color=RGBColor(31, 78, 120))
format_run(intro.add_run("Enter the expected monthly usage and vendor unit price for each line. Mark only genuinely unavoidable services as Essential. Use the summary at the end to distinguish bare-minimum survival from the current baseline and a growth-ready plan. All monetary fields are intentionally blank because pricing depends on account tier, region, volume, and negotiated terms."), size=9.5)

doc.add_heading("Planning assumptions", level=1)
assumptions = doc.add_table(rows=4, cols=4)
assumptions.style = "Table Grid"
assumption_rows = [
    ("Forecast currency", "____________", "Forecast period", "____________ to ____________"),
    ("Starting cash available", "$____________", "Target runway", "____________ months"),
    ("Expected active users", "____________", "Expected monthly growth", "____________ %"),
    ("Contingency rate", "____________ %", "Last pricing review", "____________"),
]
for r, values in enumerate(assumption_rows):
    for c, value in enumerate(values):
        set_cell_text(assumptions.rows[r].cells[c], value, size=9, bold=c in (0, 2))
        if c in (0, 2):
            shade(assumptions.rows[r].cells[c], LIGHT_BLUE)
set_fixed_table(assumptions, [2100, 2600, 2100, 5500])

core = [
    ("Vercel", "Hosting tier, bandwidth, function execution, and team seats", "Yes", "Recommended deployment platform; verify actual production plan."),
    ("Neon PostgreSQL", "Compute hours, storage, data transfer, and backup tier", "Yes", "Documented as the current hosted PostgreSQL setup."),
    ("Prisma", "ORM is open source; include paid platform costs only if adopted", "Yes", "Schema, migrations, client generation, and database access."),
    ("AWS S3", "GB stored + requests + internet data transfer", "Yes", "Object storage for uploads and screenshots; local upload fallback exists."),
    ("NextAuth.js", "Open source; budget hosting/database usage rather than a license", "Yes", "Session and authentication orchestration."),
    ("Google OAuth", "Usually no direct per-login fee; record any related cloud costs", "Yes", "Configured sign-in provider."),
    ("Stripe", "% + fixed fee per transaction; subscriptions and payment intents", "Yes", "Checkout, subscriptions, webhooks, and platform payments."),
    ("Stripe Connect Express", "Connected-account and payout/transfer fees", "Yes", "Creator onboarding and payout management."),
    ("Pusher", "Connections, channels, and messages/events", "Conditional", "Real-time messaging; application can fall back to polling."),
    ("Resend", "Emails sent and plan allowance", "Yes", "OTP and creator-outreach email delivery."),
    ("SMTP provider", "Mailbox/relay plan and messages sent", "Conditional", "Nodemailer-powered contact form; may be consolidated with Resend."),
]
add_section_table(doc, "1. Core production infrastructure", "Start here when calculating the minimum amount required to keep the product operational.", core)

ai = [
    ("OpenAI API", "Input/output tokens, images generated, and vision usage", "Conditional", "Used for assistant chat, translations, images, and screenshot analysis."),
    ("Anthropic API", "Input/output tokens and prompt caching", "Optional", "Claude assistant models."),
    ("DeepSeek API", "Input/output tokens", "Optional", "OpenAI-compatible assistant models."),
    ("Moonshot / Kimi API", "Input/output tokens", "Optional", "Kimi assistant models."),
    ("Google Cloud Translation", "Characters translated", "Conditional", "Content translation; application has a dictionary fallback."),
    ("Google Gemini / Veo", "Generated video seconds and resolution", "Experimental", "Used only by the video-creator test subproject."),
]
add_section_table(doc, "2. AI and translation services", "Use realistic monthly caps. For survival planning, disable or rate-limit providers that are not required for the core customer journey.", ai)

discovery = [
    ("KOL FastAPI sidecar", "Hosting compute, memory, requests, logs, and egress", "Conditional", "External creator-discovery service referenced through KOL_API_URL."),
    ("Influencers Club API", "API plan, profiles, enrichment records, or credits", "Optional", "Creator discovery and enrichment."),
    ("YouTube Data API", "Quota allocation and any associated Google Cloud costs", "Optional", "Credential is configured, but this repository contains no direct usage."),
    ("Facebook OAuth", "Usually no direct fee; include developer/platform overhead", "Optional", "Implemented, but local credentials are empty."),
    ("WeChat Open Platform", "Registration, verification, or platform service costs", "Optional", "Configured but feature-flagged and described as in development."),
]
add_section_table(doc, "3. Discovery, social, and optional integrations", "Keep these outside the bare-minimum budget unless they directly drive acquisition or contracted customer functionality.", discovery)

support = [
    ("Domain registration", "Domains × annual renewal price", "Yes", "Not represented in code, but normally required for production availability."),
    ("DNS / CDN / edge security", "Zone, traffic, requests, WAF, and protection tier", "Conditional", "May be included with Vercel or a separate provider."),
    ("Monitoring and error tracking", "Events, errors, traces, logs, and retention", "Recommended", "No dedicated provider is present in the repository."),
    ("Backups and recovery", "Storage, snapshots, retention, and restore testing", "Yes", "Confirm what Neon and S3 plans include before adding a separate line."),
    ("Source control / CI", "Seats, build minutes, artifacts, and runners", "Recommended", "GitHub is referenced in deployment instructions; no CI workflow is present."),
    ("Security and secrets management", "Seats, stored secrets, scans, and audit features", "Recommended", "Add only if not included in the hosting platform."),
    ("Support / maintenance reserve", "Engineering hours or monthly retainer", "Recommended", "Operational labor is often larger than vendor infrastructure cost."),
    ("Taxes and foreign-exchange buffer", "Percentage of vendor spend", "Recommended", "Use when vendors bill in another currency or exclude applicable taxes."),
]
add_section_table(doc, "4. Often-missed operating costs", "These are not all visible in application code, but omitting them can make a survival estimate materially too low.", support)

doc.add_heading("5. Survival-budget summary", level=1)
p = doc.add_paragraph("Transfer the totals from the detailed tables into the appropriate layer. Avoid counting the same service in more than one layer when calculating the final funding target.")
add_summary_table(doc)

doc.add_heading("Runway formulas", level=2)
formula_rows = [
    ("Annual infrastructure total", "Sum of all included annualized forecasts"),
    ("Average monthly burn", "Annual infrastructure total ÷ 12"),
    ("Cash runway", "Cash available ÷ average monthly burn"),
    ("Survival funding target", "Bare-minimum annual total + contingency reserve + one-time setup costs"),
    ("Suggested contingency", "10%–20% × annual operating infrastructure, adjusted for usage uncertainty"),
]
ft = doc.add_table(rows=1, cols=2)
ft.style = "Table Grid"
set_cell_text(ft.rows[0].cells[0], "Metric", bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(ft.rows[0].cells[1], "Formula", bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
shade(ft.rows[0].cells[0], BLUE)
shade(ft.rows[0].cells[1], BLUE)
for idx, values in enumerate(formula_rows):
    cells = ft.add_row().cells
    set_cell_text(cells[0], values[0], size=9, bold=True)
    set_cell_text(cells[1], values[1], size=9)
    if idx % 2 == 1:
        shade(cells[0], PALE_BLUE)
        shade(cells[1], PALE_BLUE)
set_fixed_table(ft, [3300, 8900])

doc.add_heading("Final review checklist", level=2)
checks = [
    "☐ Every included service has a documented pricing source and review date.",
    "☐ Usage assumptions reflect a low, expected, and high scenario where costs are variable.",
    "☐ Free-tier limits and overage rates have both been considered.",
    "☐ Transactional costs such as Stripe fees scale with forecast payment volume.",
    "☐ AI, storage, bandwidth, email, and real-time messaging have explicit monthly caps or alerts.",
    "☐ Taxes, currency conversion, backups, monitoring, domain renewal, and maintenance are included.",
    "☐ The bare-minimum plan has a documented list of features that would be disabled during survival mode.",
]
for check in checks:
    p = doc.add_paragraph(check)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(3)

doc.core_properties.title = "Overseed Infrastructure Survival Budget"
doc.core_properties.subject = "Editable infrastructure pricing and annual forecast worksheet"
doc.core_properties.author = "Overseed"
doc.save(OUT)
print(OUT)
